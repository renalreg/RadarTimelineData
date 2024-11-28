import inspect
import os
import random
from dataclasses import dataclass
from typing import Any, List

import docx
import polars as pl
import xlsxwriter
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import (
    WD_COLOR_INDEX,
    WD_ALIGN_PARAGRAPH,
)
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm, Mm
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from loguru import logger
from openpyxl.utils import get_column_letter

from radar_timeline_data.audit_writer.stylesheet import (
    stylesheet,
    table_styles,
    page_color,
)


@dataclass
class Table:
    table_name: str
    table: pl.DataFrame
    text: str


@dataclass
class ComparisonTable:
    table_sheet: str
    old_tables: list[Table]
    new_table: Table
    common_keys: list[str]


@dataclass
class Heading:
    text: str
    style: str


@dataclass
class List:
    text: str | None | Heading
    elements: list


@dataclass
class Change:
    description: str
    changes: list


@dataclass
class WorkSheet:
    name: str


class AuditWriter:
    """
    A class to manage the creation and writing of audit documents and spreadsheets.

    This class is responsible for initializing an audit document and spreadsheet,
    adding headings and paragraphs to the document, and managing the creation of
    worksheets within the spreadsheet.

    Attributes:
    - directory (str): The directory where the output files will be saved.
    - filename (str): The name of the audit file.
    - document (Document): An instance of the Document class for creating the audit document.
    - important_High (int): A counter for high importance audit entries.
    - important_Low (int): A counter for low importance audit entries.
    - info (dict): A dictionary to store additional information related to the audit.
    - wb (Workbook): An instance of the Workbook class for creating the audit spreadsheet.
    - current_worksheet (Worksheet): The current worksheet being written to in the spreadsheet.
    - worksheets (dict): A dictionary of worksheets within the spreadsheet.
    """

    def __init__(
        self,
        directory: str,
        filename: str,
        document_title: str,
        include_excel: bool = True,
        include_breakdown: bool = True,
        include_logger: bool = True,
    ):
        """
        Initializes an AuditWriter object.

        Parameters:
        - directory (str): The directory where the output files will be saved.
        - filename (str): The name of the audit file.
        - include_excel (bool): Whether to allow excel files to be generated
        - include_breakdown (bool): Whether to include top info section in doc contain info such as warnings
        """

        # file init
        self.directory = directory
        self.filename = filename
        self.document = Document()
        # Access the first section of the document
        section = self.document.sections[0]

        # Set page size to A4
        section.page_height = Mm(297)
        section.page_width = Mm(210)

        self.add_watermark()

        # set style
        self.stylesheet = stylesheet
        self.__style()

        # add heading
        para = self.document.add_heading(f"Audit {document_title}", 0)
        self.add_paragraph_border(para, ["bottom"])
        # add process sub heading
        para = self.document.add_paragraph("Proccess", style="Heading 1")
        self.add_paragraph_border(para, ["bottom"])
        # for top breakdown
        self.__include_breakdown = include_breakdown
        if include_breakdown:
            self.important_High = 0
            self.important_Low = 0
            self.info = {}

        # for excel
        self.__include_excel = include_excel
        if include_excel:
            self.wb = xlsxwriter.Workbook(
                os.path.join(self.directory, f"{self.filename}.xlsx")
            )
            self.current_worksheet = None
            self.worksheets = {}

        # select logger object
        self.__logger = logger if include_logger else StubObject()

    def __style(self):
        """
        Applies the styles from the stylesheet to the document.
        """
        for style_name, style_attributes in self.stylesheet.items():
            if style_name in self.document.styles:
                style = self.document.styles[style_name]
            elif style_name in ["Symbol"]:
                style = self.document.styles.add_style(
                    style_name, docx.enum.style.WD_STYLE_TYPE.CHARACTER
                )
            else:
                style = self.document.styles.add_style(
                    style_name, docx.enum.style.WD_STYLE_TYPE.PARAGRAPH
                )

            font = style.font
            if style_name == "Title" or "Heading" in style_name:
                style.element.rPr.rFonts.set(
                    qn("w:asciiTheme"), style_attributes.get("font", "Times New Roman")
                )
            else:
                font.name = style_attributes.get("font", "Times New Roman")
            font.size = style_attributes.get("size", Pt(12))
            font.bold = style_attributes.get("bold", False)
            font.color.rgb = style_attributes.get("color", RGBColor(0, 0, 0))

            if style_name not in ["Symbol"]:
                paragraph_format = style.paragraph_format
                paragraph_format.alignment = style_attributes.get(
                    "alignment", WD_ALIGN_PARAGRAPH.LEFT
                )

    def add_watermark(self, watermark_path: str = None):
        if not watermark_path:
            script_dir = os.path.dirname(__file__)
            # Construct the full path to the file
            print(script_dir)
            watermark_path = os.path.join(script_dir, "img/watermark.png")
            print(watermark_path)

        # Open the document footer for editing
        section = self.document.sections[0]
        footer = section.footer

        # Create a paragraph and add image to it
        paragraph = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )
        run = paragraph.add_run()
        run.add_picture(watermark_path, width=Inches(1))  # Adjust width as necessary

        # Center align the paragraph
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Set the image as behind text
        shape = run._element.getparent()
        shape._inline = None  # Make the shape floating
        shape.wrap = True
        shape.width = Inches(1)  # Adjust width as necessary
        shape.height = Inches(1)  # Adjust height as necessary

    def add_change(self, changes: list[Any]):
        # description of the change

        total_length = 0
        for change in changes:
            if isinstance(change, list):
                for item in change:
                    if isinstance(item, str):
                        total_length += len(item)
                    else:
                        return
            elif isinstance(change, str):
                total_length += len(change)
            elif isinstance(change, pl.DataFrame):
                total_length += sum(len(col) for col in change.columns)
            else:
                return

        para = self.document.add_paragraph()
        # self.__set_paragraph_spacing(para, 0, 0)

        if total_length < 60:
            for index, change in enumerate(changes):
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if index != 0:
                    run = para.add_run(" \u2192 ")
                    run.style = self.document.styles["Symbol"]
                if isinstance(change, pl.DataFrame):
                    para.add_run(f"{str(change.columns)} ")
                else:
                    para.add_run(f"{str(change)} ")
        else:
            self.add_change_table(changes)
            self.document.add_paragraph("\n")

    def add_hyperlink(self, paragraph, url, text):
        """
        Adds a hyperlink to the audit document.

        Parameters:
        - paragraph: The paragraph to which the hyperlink will be added.
        - url (str): The URL of the hyperlink.
        - text (str): The text to be displayed for the hyperlink.
        """

        part = paragraph.part
        r_id = part.relate_to(
            url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        hyperlink = docx.oxml.shared.OxmlElement("w:hyperlink")
        hyperlink.set(docx.oxml.shared.qn("r:id"), r_id)

        new_run = docx.oxml.shared.OxmlElement("w:r")
        rPr = docx.oxml.shared.OxmlElement("w:rPr")
        new_run.append(rPr)
        new_run.text = text.replace("_", " ")
        hyperlink.append(new_run)

        r = paragraph.add_run()
        r._r.append(hyperlink)

        link = self.stylesheet.get("Link")
        color = link.get("color") if link is not None else None
        r.font.color.rgb = (
            self.stylesheet.get("Link").get("color") if color else RGBColor(255, 0, 0)
        )
        r.font.underline = True

        return hyperlink

    def add_important(self, text: str, severity: bool):
        """
        Adds important information to the audit document.

        Parameters:
        - text (str): The text of the important information.
        - severity (bool): The severity level of the information (True for high severity, False for low severity).
        """
        paragraph = self.document.add_paragraph()
        self.__set_paragraph_spacing(paragraph, 0, 0)
        run = paragraph.add_run("\u26A0")
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
        self.__format_run(run, Pt(16), (204, 51, 0)) if severity else self.__format_run(
            run, Pt(16), (255, 204, 0)
        )
        run = paragraph.add_run(text)
        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
        if severity:
            self.important_High += 1
            self.__logger.critical(text)
        else:
            self.important_Low += 1
            self.__logger.warning(text)

    def add_info(self, key: str, value: str | tuple[str, str]):
        """
        Adds information to the audit writer.

        Args:
            key (str): The key of the information to add.
            value (str): The value of the information to add.

        Returns:
            None
        """
        if type(value) is tuple:
            if key in self.info:
                self.info[key][value[0]] = value[1]
            else:
                self.info[key] = {value[0]: value[1]}
        else:
            self.info[key] = value
        self.__logger.info(f"{key} : {value}")

    def _comparison_table(self, element: ComparisonTable):
        def comparison_table(
            old_tables: pl.DataFrame, new_tables: pl.DataFrame, common_keys
        ):
            # Group the old table by common_keys and count the number of rows for each group
            result_old = old_tables.groupby(common_keys).agg(
                pl.count().alias("old_count")
            )

            # Group the new table by common_keys and count the number of rows for each group
            result_new = new_tables.groupby(common_keys).agg(
                pl.count().alias("new_count")
            )

            result = result_old.join(result_new, on=common_keys, how="outer")
            result = result.with_columns(
                pl.coalesce([i, i + "_right"]).alias(i) for i in common_keys
            ).drop(i + "_right" for i in common_keys)
            result = (
                result.fill_null(0)
                .with_columns(
                    (pl.col("old_count") - pl.col("new_count")).alias("count")
                )
                .drop("old_count", "new_count")
            )

            new_df = pl.DataFrame(schema=new_tables.schema)
            old_df = pl.DataFrame(schema=old_tables.schema)
            new_tables = new_tables.sort(common_keys)
            old_tables = old_tables.sort(common_keys)

            result = result.sort(common_keys)
            for i in result.iter_rows():
                *keys, count = i

                # Build the filter condition by combining multiple conditions with `&`
                filter_condition = (
                    pl.col(common_keys[0]) == keys[0]  # First key-value comparison
                )
                for k, v in zip(common_keys[1:], keys[1:]):
                    filter_condition = filter_condition & (pl.col(k) == v)

                # Find the rows in new_tables that match the keys
                new_table_matching_rows = new_tables.filter(filter_condition)
                old_table_matching_rows = old_tables.filter(filter_condition)
                # Insert rows based on the count value
                if count > 0:
                    # Add `count` blank rows before the matching rows
                    blank_rows = pl.DataFrame(
                        {
                            col: (
                                [keys[0]] * count
                                if col == "patient_id"
                                else [None] * count
                            )
                            for col in new_tables.columns
                        },
                        schema=new_tables.schema,
                    )
                    new_df = pl.concat([new_df, blank_rows, new_table_matching_rows])
                    old_df = pl.concat([old_df, old_table_matching_rows])

                elif count == 0:
                    # Add matching rows directly
                    new_df = pl.concat([new_df, new_table_matching_rows])
                    old_df = pl.concat([old_df, old_table_matching_rows])
                elif count < 0:
                    # Add matching rows first, followed by `-count` blank rows
                    blank_rows = pl.DataFrame(
                        {col: [None] * abs(count) for col in old_tables.columns},
                        schema=old_tables.schema,
                    )
                    old_df = pl.concat([old_df, blank_rows, old_table_matching_rows])
                    new_df = pl.concat([new_df, new_table_matching_rows])
            return new_df, old_df

        old_tables = [i.table for i in element.old_tables]
        new_tables = element.new_table
        common_keys = element.common_keys
        if isinstance(old_tables, list):
            old_tables = pl.concat(old_tables)
        new_df, old_df = comparison_table(old_tables, new_tables.table, common_keys)
        self.set_ws(element.table_sheet)
        self.add_table(
            text=element.old_tables[0].text,
            table=old_df,
            table_name=element.old_tables[0].table_name,
            indent_level=0,
        )
        self.add_table(
            text=element.new_table.text,
            table=new_df,
            table_name=element.new_table.table_name,
            indent_level=0,
        )

    def add_table(
        self, text: str, table: pl.DataFrame, table_name: str, indent_level: int = 0
    ):
        """
        Adds a table to the xlsx document and creates a link in the current document.

        Parameters:
        - text (str): Text description for the table.
        - table (pl.DataFrame): The DataFrame to be added as a table.
        - table_name (str): The name of the table must not contain spaces.
        """
        if self.__include_excel:
            # If indented, apply bullet points with "List Bullet" style
            if indent_level > 0:
                para = self.document.add_paragraph(
                    f"{text} \u2192 ", style="List Bullet"
                )
            else:
                para = self.document.add_paragraph(f"{text} \u2192 ")

            # Apply indentation based on the level
            if indent_level > 0:
                para.paragraph_format.left_indent = Inches(0.25 * indent_level)

            table_name = table_name.strip()
            self.add_hyperlink(
                para,
                f"{self.filename}.xlsx#{self.current_worksheet}!{get_column_letter(self.worksheets[self.current_worksheet] + 1)}4",
                table_name,
            )

            random_style = random.choice(table_styles)
            table.write_excel(
                workbook=self.wb,
                worksheet=self.current_worksheet,
                table_name=table_name,
                table_style=random_style,
                position=(3, self.worksheets[self.current_worksheet]),
                include_header=True,
            )

            name_format = self.wb.add_format({"bold": True, "font_size": 18})

            self.wb.get_worksheet_by_name(self.current_worksheet).write(
                f"{get_column_letter(self.worksheets[self.current_worksheet] + 1)}2",
                table_name.replace("_", " "),
                name_format,
            )

            self.worksheets[self.current_worksheet] += len(table.columns) + 1
            call = inspect.getframeinfo(inspect.currentframe().f_back)
            self.__logger.info(
                f"{call.function}:{call.lineno} : {table_name} created "
                f"at file path {self.filename}.xlsx#{self.current_worksheet}!"
                f"{get_column_letter(self.worksheets[self.current_worksheet] + 1)}4"
            )

    def add_table_snippets(self, table: pl.DataFrame):
        """
        Adds a Python-docx table with column names and types, as well as some rows, based on a polars DataFrame.

        Args:
            table (polars.DataFrame): The DataFrame containing the data.

        Returns:
            None
        """
        cols = table.columns
        doc_tbl = self.document.add_table(rows=1, cols=len(cols))
        doc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc_tbl.style = "Table Grid"
        doc_tbl.autofit = True
        hdr_cells = doc_tbl.rows[0].cells
        for index, (name, data_type) in enumerate(zip(cols, table.dtypes)):
            hdr_cells[index].text = name + "\n" + str(data_type)

    def add_change_table(self, changes: list[Any]):
        if not 0 < len(changes) < 8:
            return
        doc_tbl = self.document.add_table(rows=1, cols=((len(changes) * 2) - 1))
        doc_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        doc_tbl.style = "Table Grid"
        doc_tbl.autofit = True

        tblBorders = parse_xml(
            r'<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            r'<w:top w:val="nil"/><w:left w:val="nil"/><w:bottom w:val="nil"/><w:right w:val="nil"/>'
            r'<w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
            r"</w:tblBorders>"
        )
        doc_tbl._tbl.tblPr.append(tblBorders)
        hdr_cells = doc_tbl.rows[0].cells

        for index, change in enumerate(changes):
            max_length = 1
            if index != 0:  # For every iteration apart from the first
                arrow_cell = hdr_cells[((index * 2) - 1)]
                arrow_cell.width = Cm(1.0)
                arrow_cell.text = "\u2192"
                arrow_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                for paragraph in arrow_cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.style = "Symbol"

            cell = hdr_cells[(index * 2) if index != 0 else 0]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if isinstance(change, str):
                cell.text = change
                max_length = len(change)
            elif isinstance(change, list) and all(
                isinstance(item, str) for item in change
            ):
                cell.text = "\n".join(change)
                max_length = len(max(change, key=len))
            elif isinstance(change, pl.DataFrame):
                temp = change.columns
                cell.text = "\n".join(temp)
                max_length = len(max(temp, key=len))

            cell.width = Cm(min(round(0.42 * max_length), 5))  # Default larger width

    def add_text(self, text: str, style: str | None = None, indent_level: int = 0):
        """
        Adds text to the audit document.

        Parameters:
        - text (str): The text to be added.
        - style (str | None): The style to apply to the text.
        - indent_level (int): The level of indentation, where each level adds more indentation.
        """
        if indent_level > 0:
            # Add a paragraph with bullet point style for indented text
            para = self.document.add_paragraph(text, style="List Bullet")
        else:
            # Regular paragraph for non-indented text
            para = self.document.add_paragraph(text)

        if style:
            para.style = style

        # Set the indent based on the level of indentation
        if indent_level > 0:
            para.paragraph_format.left_indent = Inches(0.25 * indent_level)

        self.__logger.info(text)

    def add(self, element: Any, indent_level: int = 0):
        """
        Adds an element (text, list, table) to the document.

        Parameters:
        - element (Any): The element to be added.
        - indent_level (int): The level of indentation to be applied. Nested lists will increase this.
        """
        if isinstance(element, str):
            self.add_text(element, indent_level=indent_level)

        if isinstance(element, list):
            for i in element:
                self.add(i)

        if isinstance(element, Table):
            element: Table
            self.add_table(
                text=element.text,
                table=element.table,
                table_name=element.table_name,
                indent_level=indent_level,
            )

        if isinstance(element, ComparisonTable):
            element: ComparisonTable
            self._comparison_table(element)
        if isinstance(element, List):
            element: List
            if element.text:
                if isinstance(element.text, str):
                    self.add_text(element.text, indent_level=indent_level)
                elif isinstance(element.text, Heading):
                    self.add_text(
                        element.text.text,
                        style=element.text.style,
                        indent_level=indent_level,
                    )
            for i in element.elements:
                # Recursively add elements with increased indentation for nested elements
                self.add(i, indent_level=indent_level + 1)
        if isinstance(element, Heading):
            element: Heading
            self.add_text(element.text, style=element.style, indent_level=indent_level)

        if isinstance(element, Change):
            element: Change
            self.add_text(element.description, indent_level=indent_level)
            self.add_change(element.changes)

        if isinstance(element, WorkSheet):
            element: WorkSheet
            self.set_ws(element.name)

    def add_top_breakdown(self):
        """
        Adds a breakdown of important information at the top of the audit document.
        """
        if not self.__include_breakdown:
            return

        paragraph = self.document.paragraphs[1].insert_paragraph_before()
        run = paragraph.add_run("\u26A0")
        self.__format_run(run, Pt(16), (255, 204, 0))
        paragraph.add_run(f" {str(self.important_Low)} Warnings raised \n")

        run = paragraph.add_run("\u26A0")
        self.__format_run(run, Pt(16), (204, 51, 0))
        paragraph.add_run(f" {str(self.important_High)} Issues raised \n")

        for element in self.info:
            if type(self.info[element]) is dict:
                paragraph.add_run(f"{element}:\n")
                for key, value in self.info[element].items():
                    run = paragraph.add_run(f"\t• {key}: {value}\n")
            else:
                paragraph.add_run(f"{element}: {self.info[element]}\n")

        paragraph.insert_paragraph_before("breakdown", style="Heading 1")

    @staticmethod
    def __format_run(run: Run, font_size: int, color_rgb: tuple):
        """
        Formats the given run with specified font size and color.
        """
        run.font.size = font_size
        run.font.color.rgb = RGBColor(*color_rgb)

    @staticmethod
    def __set_paragraph_spacing(
        paragraph: Paragraph, space_before: int, space_after: int
    ):
        """
        Sets the spacing for the given paragraph.
        """
        paragraph.paragraph_format.space_before = space_before
        paragraph.paragraph_format.space_after = space_after

    @staticmethod
    def add_paragraph_border(paragraph, border_override: list[str] = None):
        # Get the XML of the paragraph
        p = paragraph._element

        # Create a new element for paragraph borders
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            p.insert(0, pPr)

        # Create the border element
        pBdr = OxmlElement("w:pBdr")

        # Define each side of the border
        for border_position in (
            border_override if border_override else ["top", "bottom"]
        ):
            border = OxmlElement(f"w:{border_position}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "10")  # Size of the border
            border.set(qn("w:space"), "1")
            border.set(qn("w:color"), "000000")  # Black color
            pBdr.append(border)

        # Append the border element to the paragraph properties
        pPr.append(pBdr)

    def set_page_color(self):
        """
        Set the background color of all pages in a Word document.

        Parameters:
        - doc: Document object from python-docx.
        - color: RGBColor object representing the color.
        """
        # Iterate through all sections in the document
        shd = OxmlElement("w:background")
        # Add attributes to the xml element
        shd.set(qn("w:color"), page_color)
        # Add background element at the start of Document.xml using below
        self.document.element.insert(0, shd)
        background_shp = OxmlElement(
            "w:displayBackgroundShape"
        )  # Setting to use my background
        self.document.settings.element.insert(0, background_shp)  # Apply setting

    def commit_audit(self):
        """
        Commits the audit by adding the top breakdown, closing the workbook, and saving the document.
        """
        self.add_top_breakdown()
        self.set_page_color()

        self.wb.close()
        self.document.save(os.path.join(self.directory, f"{self.filename}.docx"))

    def set_ws(self, worksheet_name: str):
        """
        sets a worksheet to the audit file. can be used to change worksheets

        Args:
            worksheet_name (str): The name of the worksheet.
        """
        if worksheet_name not in self.worksheets:
            self.worksheets[worksheet_name] = 0
        self.current_worksheet = worksheet_name


class StubObject:
    def __init__(self):
        self.info = []

    def __getattr__(self, name):
        return self._stub_callable

    def __setattr__(self, name, value):
        if name == "info" and not hasattr(self, "info"):
            object.__setattr__(self, "info", [])
        else:
            pass

    def _stub_callable(self, *args, **kwargs):
        return self
